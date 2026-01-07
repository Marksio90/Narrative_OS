"""
DOCX Export Generator
Professional manuscript formatting for Word documents
Uses python-docx for modern .docx generation
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from typing import Optional, List, Dict, Any
from datetime import datetime
import io


class DocxGenerator:
    """
    Generate professional DOCX manuscripts

    Features:
    - Cover page with title, author, metadata
    - Table of contents (bookmarks)
    - Chapter formatting with page breaks
    - Scene separators (### or custom)
    - Professional typography (Times New Roman, 12pt)
    - Standard manuscript margins (1" all sides)
    """

    def __init__(self):
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """Configure professional manuscript styles"""
        # Page setup - standard manuscript format
        section = self.doc.sections[0]
        section.page_height = Inches(11)  # Letter size
        section.page_width = Inches(8.5)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        # Title style
        title_style = self.doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_format = title_style.font
        title_format.name = 'Times New Roman'
        title_format.size = Pt(18)
        title_format.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(24)

        # Chapter heading style
        chapter_style = self.doc.styles.add_style('ChapterHeading', WD_STYLE_TYPE.PARAGRAPH)
        chapter_format = chapter_style.font
        chapter_format.name = 'Times New Roman'
        chapter_format.size = Pt(14)
        chapter_format.bold = True
        chapter_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        chapter_style.paragraph_format.space_before = Pt(0)
        chapter_style.paragraph_format.space_after = Pt(12)

        # Body text style (standard manuscript format)
        body_style = self.doc.styles['Normal']
        body_format = body_style.font
        body_format.name = 'Times New Roman'
        body_format.size = Pt(12)
        body_style.paragraph_format.line_spacing = 2.0  # Double-spaced
        body_style.paragraph_format.first_line_indent = Inches(0.5)
        body_style.paragraph_format.space_after = Pt(0)

    def add_cover_page(
        self,
        title: str,
        author: str,
        genre: Optional[str] = None,
        word_count: Optional[int] = None,
        subtitle: Optional[str] = None
    ):
        """
        Add professional cover page
        Standard manuscript format with centered title block
        """
        # Add vertical space (approx 1/3 page)
        for _ in range(8):
            self.doc.add_paragraph()

        # Title
        p = self.doc.add_paragraph(title, style='CustomTitle')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle
        if subtitle:
            p = self.doc.add_paragraph(subtitle)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.italic = True

        # Spacing
        self.doc.add_paragraph()

        # Author
        p = self.doc.add_paragraph(f'by {author}')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

        # Metadata (bottom of page)
        for _ in range(10):
            self.doc.add_paragraph()

        # Genre and word count
        if genre or word_count:
            metadata_lines = []
            if genre:
                metadata_lines.append(f'Genre: {genre}')
            if word_count:
                metadata_lines.append(f'Word Count: {word_count:,}')

            p = self.doc.add_paragraph(' | '.join(metadata_lines))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(100, 100, 100)

        # Page break
        self.doc.add_page_break()

    def add_chapter(
        self,
        title: str,
        content: str,
        chapter_number: Optional[int] = None
    ):
        """
        Add a chapter with proper formatting

        Args:
            title: Chapter title (e.g., "The Beginning")
            content: Chapter prose text
            chapter_number: Optional chapter number (e.g., 1, 2, 3)
        """
        # Chapter heading
        if chapter_number:
            heading_text = f'CHAPTER {chapter_number}\n{title}'
        else:
            heading_text = title

        p = self.doc.add_paragraph(heading_text, style='ChapterHeading')

        # Process content - handle scene breaks
        scenes = content.split('\n\n###\n\n')  # Scene separator

        for i, scene in enumerate(scenes):
            # Add scene separator (except for first scene)
            if i > 0:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run('###')
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                self.doc.add_paragraph()  # Blank line after separator

            # Add scene paragraphs
            paragraphs = scene.strip().split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    self.doc.add_paragraph(para_text.strip())

        # Page break after chapter
        self.doc.add_page_break()

    def add_table_of_contents(self, chapters: List[Dict[str, Any]]):
        """
        Add simple table of contents

        Args:
            chapters: List of dicts with 'number' and 'title' keys
        """
        p = self.doc.add_paragraph('TABLE OF CONTENTS', style='ChapterHeading')

        self.doc.add_paragraph()

        for chapter in chapters:
            number = chapter.get('number', '')
            title = chapter.get('title', 'Untitled')

            p = self.doc.add_paragraph()
            run = p.add_run(f'Chapter {number}: {title}')
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            p.paragraph_format.left_indent = Inches(0.5)

        self.doc.add_page_break()

    def add_front_matter(self, content: str, title: str = "Prologue"):
        """Add prologue, foreword, or other front matter"""
        p = self.doc.add_paragraph(title, style='ChapterHeading')

        paragraphs = content.strip().split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                self.doc.add_paragraph(para_text.strip())

        self.doc.add_page_break()

    def add_back_matter(self, content: str, title: str = "Epilogue"):
        """Add epilogue, afterword, or other back matter"""
        self.add_front_matter(content, title)

    def save(self, output_path: str):
        """Save document to file"""
        self.doc.save(output_path)

    def to_bytes(self) -> bytes:
        """Return document as bytes for download"""
        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()


def generate_manuscript_docx(
    project_title: str,
    author_name: str,
    chapters: List[Dict[str, Any]],
    genre: Optional[str] = None,
    prologue: Optional[str] = None,
    epilogue: Optional[str] = None,
    include_toc: bool = True
) -> bytes:
    """
    Generate complete manuscript DOCX

    Args:
        project_title: Book title
        author_name: Author name
        chapters: List of chapter dicts with 'title', 'content', 'number'
        genre: Optional genre
        prologue: Optional prologue text
        epilogue: Optional epilogue text
        include_toc: Whether to include table of contents

    Returns:
        DOCX file as bytes
    """
    generator = DocxGenerator()

    # Calculate word count
    total_words = 0
    for chapter in chapters:
        content = chapter.get('content', '')
        total_words += len(content.split())

    # Cover page
    generator.add_cover_page(
        title=project_title,
        author=author_name,
        genre=genre,
        word_count=total_words
    )

    # Table of contents
    if include_toc and len(chapters) > 1:
        toc_data = [
            {'number': ch.get('number', i+1), 'title': ch.get('title', 'Untitled')}
            for i, ch in enumerate(chapters)
        ]
        generator.add_table_of_contents(toc_data)

    # Prologue
    if prologue:
        generator.add_front_matter(prologue, 'Prologue')

    # Chapters
    for i, chapter in enumerate(chapters):
        generator.add_chapter(
            title=chapter.get('title', 'Untitled'),
            content=chapter.get('content', ''),
            chapter_number=chapter.get('number', i + 1)
        )

    # Epilogue
    if epilogue:
        generator.add_back_matter(epilogue, 'Epilogue')

    return generator.to_bytes()
